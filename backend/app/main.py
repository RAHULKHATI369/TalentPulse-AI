import json
import os
import asyncio
from fastapi import FastAPI, Depends, HTTPException, BackgroundTasks, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from fastapi.staticfiles import StaticFiles
from sqlalchemy.orm import Session
from typing import List, Dict

from backend.app.config import settings
from backend.app.database import engine, Base, get_db
from backend.app.models import Candidate, InterviewLog
from backend.app import schemas
from backend.app import gemini_service
from backend.app import anomaly_detector
from backend.app import simulator
from backend.app import hooks

# Create FastAPI app instance
app = FastAPI(
    title="TalentPulse AI - Core API Engine",
    description="Enterprise-grade candidate screening and behavioral simulator engine.",
    version="1.0.0"
)

# Configure CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Ensure database tables are created on startup
@app.on_event("startup")
def startup_event():
    Base.metadata.create_all(bind=engine)

# ==========================================
# CANDIDATE MANAGEMENT ENDPOINTS
# ==========================================

async def process_candidate_background(candidate_id: int, job_description: str = None):
    """
    Background worker task to fetch Gemini attributes and calculate embeddings.
    """
    # Create a new local DB session for background task safety
    from backend.app.database import SessionLocal
    db = SessionLocal()
    try:
        candidate = db.query(Candidate).filter(Candidate.id == candidate_id).first()
        if not candidate:
            return

        # 1. Generate Tone / Persona Mapping
        persona = await gemini_service.generate_candidate_tone_and_persona(
            name=candidate.name,
            skills=candidate.skills,
            resume_text=candidate.resume_text
        )
        candidate.resume_tone = persona

        # 2. Generate Candidate Embedding (representing skills and experience)
        vector_text = f"Candidate Name: {candidate.name}. Skills: {candidate.skills}. Experience: {candidate.experience_years} years. Narrative: {candidate.resume_text or ''}"
        vector = await gemini_service.get_embedding_vector(vector_text)
        candidate.embedding = json.dumps(vector)

        # 3. Generate Why Report
        why_report = await gemini_service.generate_why_this_candidate_report(
            name=candidate.name,
            skills=candidate.skills,
            experience_years=candidate.experience_years,
            persona=persona,
            job_desc=job_description
        )
        candidate.why_report = why_report

        db.commit()

        # Recalculate Z-scores for database context
        anomaly_detector.recalculate_z_scores(db)
        
        # Reload candidate to obtain Z-Score details
        db.refresh(candidate)
        
        # 4. Trigger external webhook validation loop (Hugging Face ml-intern)
        await hooks.trigger_ml_intern_validation_hook(
            candidate_id=candidate.id,
            candidate_name=candidate.name,
            z_score=candidate.z_score or 0.0,
            is_anomaly=candidate.is_anomaly or False
        )

    except Exception as e:
        gemini_service.logger.error(f"Error processing candidate in background: {e}")
    finally:
        db.close()

@app.post("/api/candidates", response_model=schemas.CandidateResponse)
async def create_candidate(
    candidate: schemas.CandidateCreate,
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db)
):
    """
    Creates a candidate in database, triggers Z-Score recalculation, 
    and launches async Gemini parsing pipelines.
    """
    db_candidate = db.query(Candidate).filter(Candidate.email == candidate.email).first()
    if db_candidate:
        raise HTTPException(status_code=400, detail="Candidate email already registered.")

    new_candidate = Candidate(
        name=candidate.name,
        email=candidate.email,
        experience_years=candidate.experience_years,
        skills=candidate.skills,
        resume_text=candidate.resume_text
    )
    db.add(new_candidate)
    db.commit()
    db.refresh(new_candidate)

    # Queue background AI tasks
    background_tasks.add_task(process_candidate_background, new_candidate.id)
    
    return new_candidate

@app.get("/api/candidates", response_model=List[schemas.CandidateResponse])
def get_candidates(db: Session = Depends(get_db)):
    """
    Fetches all candidates.
    """
    return db.query(Candidate).order_by(Candidate.match_score.desc(), Candidate.created_at.desc()).all()

@app.delete("/api/candidates/{candidate_id}")
def delete_candidate(candidate_id: int, db: Session = Depends(get_db)):
    """
    Deletes a candidate by ID.
    """
    candidate = db.query(Candidate).filter(Candidate.id == candidate_id).first()
    if not candidate:
        raise HTTPException(status_code=404, detail="Candidate not found.")
    db.delete(candidate)
    db.commit()
    return {"message": "Candidate deleted successfully"}

# ==========================================
# SEED DATA / BULK UPLOAD ENDPOINT
# ==========================================

@app.post("/api/candidates/upload-seed")
async def seed_candidates(db: Session = Depends(get_db)):
    """
    Bulk uploads all candidates defined in seed_data.json and evaluates them.
    This simulates file dropping ingestion with progress tracker handlers.
    """
    seed_file_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), "seed_data.json")
    if not os.path.exists(seed_file_path):
        raise HTTPException(status_code=404, detail="Seed data JSON file not found.")

    with open(seed_file_path, "r") as f:
        candidates_data = json.load(f)

    uploaded_count = 0
    skipped_count = 0
    processing_tasks = []

    for c in candidates_data:
        # Avoid duplicate email entries
        existing = db.query(Candidate).filter(Candidate.email == c["email"]).first()
        if existing:
            skipped_count += 1
            continue

        new_cand = Candidate(
            name=c["name"],
            email=c["email"],
            experience_years=c["experience_years"],
            skills=c["skills"],
            resume_text=c["resume_text"]
        )
        db.add(new_cand)
        db.commit()
        db.refresh(new_cand)
        uploaded_count += 1

        # Push to async tasks processing immediately (simulated concurrency pool)
        processing_tasks.append(process_candidate_background(new_cand.id))

    # Process async tasks to finish API seeding block in one call
    if processing_tasks:
        await asyncio.gather(*processing_tasks)

    # Recalculate Z scores for the group
    anomaly_detector.recalculate_z_scores(db)

    return {
        "status": "success",
        "inserted": uploaded_count,
        "skipped": skipped_count,
        "total": len(candidates_data)
    }

def extract_pdf_text_bytes(file_bytes: bytes) -> str:
    """
    Parses a binary PDF using pypdf and extracts raw text string.
    """
    import io
    from pypdf import PdfReader
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        text = ""
        for page in reader.pages:
            t = page.extract_text()
            if t:
                text += t + "\n"
        return text.strip()
    except Exception as e:
        raise ValueError(f"Failed to parse PDF bytes: {e}")

def extract_docx_text_bytes(file_bytes: bytes) -> str:
    """
    Parses a Microsoft Word Document (.docx) using python-docx and extracts text.
    """
    import docx
    import io
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        paragraphs = [para.text for para in doc.paragraphs if para.text]
        
        # Read text from tables inside docx as well!
        table_text = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        table_text.append(cell.text.strip())
                        
        full_text = "\n".join(paragraphs + table_text)
        return full_text.strip()
    except Exception as e:
        raise ValueError(f"Failed to parse Word Document (.docx) bytes: {e}")

@app.post("/api/candidates/upload-pdf")
async def upload_candidates_pdf(
    files: List[UploadFile] = File(...),
    db: Session = Depends(get_db)
):
    """
    Receives lists of raw binary PDFs or DOCX files, extracts text running on
    concurrent streams throttled to 5 threads using asyncio.Semaphore.
    """
    import asyncio
    
    semaphore = asyncio.Semaphore(5)
    loop = asyncio.get_event_loop()
    
    processed_candidates = []
    errors = []

    async def process_single_file(upload_file: UploadFile):
        async with semaphore:
            try:
                # 1. Read binary content
                content = await upload_file.read()
                
                # 2. Extract text based on file format
                filename_lower = upload_file.filename.lower()
                if filename_lower.endswith(".pdf"):
                    text = await loop.run_in_executor(None, extract_pdf_text_bytes, content)
                elif filename_lower.endswith(".docx"):
                    text = await loop.run_in_executor(None, extract_docx_text_bytes, content)
                else:
                    raise ValueError(f"Unsupported file format for {upload_file.filename}")

                if not text:
                    raise ValueError(f"No text content could be extracted from {upload_file.filename}")
                
                # 3. Call Gemini to parse JSON representation and evaluations in one single call
                profile = await gemini_service.extract_structured_candidate_data(text)
                
                # 4. Save to Database
                email = profile.get("email", "unknown@example.com")
                name = profile.get("name", "Unknown Candidate")
                exp = profile.get("experience_years", 0.0)
                skills = profile.get("skills", "Technical Skills")
                persona = profile.get("persona", "Collaborator")

                # Format pros, cons, alignment reason list into why_report Markdown
                pros_list = profile.get("pros", ["Strong technical background."])
                cons_list = profile.get("cons", ["Requires standard onboarding."])
                alignment = profile.get("alignment_reason", "Aligned with standard roles.")

                pros_md = "\n".join([f"- {p}" for p in pros_list])
                cons_md = "\n".join([f"- {c}" for c in cons_list])
                why_report = (
                    f"### Evaluation Report for {name}\n\n"
                    f"**Pros**:\n{pros_md}\n\n"
                    f"**Cons**:\n{cons_md}\n\n"
                    f"**Alignment Reason**:\n- {alignment}"
                )

                # Generate embedding vector
                vector_text = f"Candidate Name: {name}. Skills: {skills}. Experience: {exp} years. Narrative: {text[:1000]}"
                vector = await gemini_service.get_embedding_vector(vector_text)

                # Locate or update candidate to keep it resilient
                existing = db.query(Candidate).filter(Candidate.email == email).first()
                if existing:
                    existing.name = name
                    existing.experience_years = exp
                    existing.skills = skills
                    existing.resume_text = text
                    existing.resume_tone = persona
                    existing.why_report = why_report
                    existing.embedding = json.dumps(vector)
                    db_cand = existing
                else:
                    db_cand = Candidate(
                        name=name,
                        email=email,
                        experience_years=exp,
                        skills=skills,
                        resume_text=text,
                        resume_tone=persona,
                        why_report=why_report,
                        embedding=json.dumps(vector)
                    )
                    db.add(db_cand)
                
                db.commit()
                db.refresh(db_cand)
                
                # Trigger Hugging Face ml-intern hook
                await hooks.trigger_ml_intern_validation_hook(
                    candidate_id=db_cand.id,
                    candidate_name=db_cand.name,
                    z_score=db_cand.z_score or 0.0,
                    is_anomaly=db_cand.is_anomaly or False
                )

                processed_candidates.append({
                    "id": db_cand.id,
                    "name": db_cand.name,
                    "email": db_cand.email,
                    "status": "Processed Successfully"
                })
            except Exception as e:
                errors.append({"filename": upload_file.filename, "error": str(e)})

    # Execute all uploads concurrently throttled by Semaphore
    tasks = [process_single_file(f) for f in files]
    await asyncio.gather(*tasks)

    # Recalculate Z-scores for database context
    anomaly_detector.recalculate_z_scores(db)

    return {
        "status": "completed",
        "processed_count": len(processed_candidates),
        "failed_count": len(errors),
        "processed": processed_candidates,
        "errors": errors
    }

@app.post("/api/candidates/upload-zip")
async def upload_candidates_zip(
    file: UploadFile = File(...),
    db: Session = Depends(get_db)
):
    """
    Ingests a single .zip file, extracts PDF and TXT documents in-memory,
    and concurrently processes them into our Gemini and database pipeline,
    throttling concurrency to 5 concurrent streams via asyncio.Semaphore.
    """
    import zipfile
    import io
    import asyncio

    if not file.filename.lower().endswith(".zip"):
        import logging
        logger = logging.getLogger("UploadZip")
        logger.error(f"ZIP ingestion rejected: filename '{file.filename}' does not end with .zip")
        raise HTTPException(status_code=400, detail="Only .zip archive files are supported.")

    try:
        zip_bytes = await file.read()
        zip_buffer = io.BytesIO(zip_bytes)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to read upload stream: {e}")

    processed_candidates = []
    errors = []
    
    semaphore = asyncio.Semaphore(5)
    loop = asyncio.get_event_loop()

    async def process_candidate_doc(filename: str, file_content: bytes):
        async with semaphore:
            try:
                # 1. Extract plain text based on file format
                if filename.lower().endswith(".pdf"):
                    text = await loop.run_in_executor(None, extract_pdf_text_bytes, file_content)
                elif filename.lower().endswith(".docx"):
                    text = await loop.run_in_executor(None, extract_docx_text_bytes, file_content)
                elif filename.lower().endswith(".txt"):
                    text = file_content.decode("utf-8", errors="ignore")
                else:
                    return # skip unsupported formats silently

                if not text or not text.strip():
                    raise ValueError("No text content could be parsed from the document.")

                # 2. Extract structured candidate JSON profile and evaluation in one single call
                profile = await gemini_service.extract_structured_candidate_data(text)

                email = profile.get("email", "unknown@example.com")
                name = profile.get("name", "Unknown Candidate")
                exp = profile.get("experience_years", 0.0)
                skills = profile.get("skills", "Technical Skills")
                persona = profile.get("persona", "Collaborator")

                # Format pros, cons, alignment reason list into why_report Markdown
                pros_list = profile.get("pros", ["Strong technical background."])
                cons_list = profile.get("cons", ["Requires standard onboarding."])
                alignment = profile.get("alignment_reason", "Aligned with standard roles.")

                pros_md = "\n".join([f"- {p}" for p in pros_list])
                cons_md = "\n".join([f"- {c}" for c in cons_list])
                why_report = (
                    f"### Evaluation Report for {name}\n\n"
                    f"**Pros**:\n{pros_md}\n\n"
                    f"**Cons**:\n{cons_md}\n\n"
                    f"**Alignment Reason**:\n- {alignment}"
                )

                # Generate embedding vector
                vector_text = f"Candidate Name: {name}. Skills: {skills}. Experience: {exp} years. Narrative: {text[:1000]}"
                vector = await gemini_service.get_embedding_vector(vector_text)

                # 3. Relational Storage committing (updates if exists to be resilient)
                existing = db.query(Candidate).filter(Candidate.email == email).first()
                if existing:
                    existing.name = name
                    existing.experience_years = exp
                    existing.skills = skills
                    existing.resume_text = text
                    existing.resume_tone = persona
                    existing.why_report = why_report
                    existing.embedding = json.dumps(vector)
                    db_cand = existing
                else:
                    db_cand = Candidate(
                        name=name,
                        email=email,
                        experience_years=exp,
                        skills=skills,
                        resume_text=text,
                        resume_tone=persona,
                        why_report=why_report,
                        embedding=json.dumps(vector)
                    )
                    db.add(db_cand)

                db.commit()
                db.refresh(db_cand)

                # Trigger Hugging Face validation webhook
                await hooks.trigger_ml_intern_validation_hook(
                    candidate_id=db_cand.id,
                    candidate_name=db_cand.name,
                    z_score=db_cand.z_score or 0.0,
                    is_anomaly=db_cand.is_anomaly or False
                )

                processed_candidates.append({
                    "id": db_cand.id,
                    "name": db_cand.name,
                    "email": db_cand.email,
                    "status": "Processed Successfully"
                })

            except Exception as e:
                errors.append({"filename": filename, "error": str(e)})

    try:
        # Open zip container
        with zipfile.ZipFile(zip_buffer) as zf:
            tasks = []
            for name in zf.namelist():
                # Skip directories and subfolders
                if name.endswith("/") or "__MACOSX" in name:
                    continue
                
                # Filter PDF, TXT, and DOCX
                if name.lower().endswith((".pdf", ".txt", ".docx")):
                    content = zf.read(name)
                    # Extract only the base file name from path
                    base_name = name.split("/")[-1]
                    tasks.append(process_candidate_doc(base_name, content))

            if not tasks:
                import logging
                logger = logging.getLogger("UploadZip")
                logger.error("ZIP ingestion failed: No valid .pdf, .txt, or .docx documents found in the archive.")
                raise HTTPException(status_code=400, detail="The .zip file does not contain any valid .pdf, .txt, or .docx documents.")

            # Process all tasks concurrently throttled by Semaphore
            await asyncio.gather(*tasks)
            
    except zipfile.BadZipFile as bz:
        import logging
        logger = logging.getLogger("UploadZip")
        logger.error(f"ZIP ingestion failed: Bad zip archive error: {bz}")
        raise HTTPException(status_code=400, detail="Uploaded file is not a valid zip archive.")

    # Recalculate Z-scores for database context
    anomaly_detector.recalculate_z_scores(db)

    return {
        "status": "completed",
        "processed_count": len(processed_candidates),
        "failed_count": len(errors),
        "processed": processed_candidates,
        "errors": errors
    }

# ==========================================
# CANDIDATE SCREENING / EMBEDDING MATCHING
# ==========================================

@app.post("/api/screen")
async def screen_candidates(job: schemas.JobDescription, db: Session = Depends(get_db)):
    """
    Performs vector similarity matches between the job requirements and all candidate embeddings.
    Updates candidate match_scores and returns top 10 ranked profiles.
    """
    # 1. Obtain embedding vector for job description using text-embedding-004
    job_embedding = await gemini_service.get_embedding_vector(f"{job.job_title} : {job.requirements}")
    
    candidates = db.query(Candidate).all()
    if not candidates:
        return {"message": "No candidates available for screening."}

    # 2. Iterate and evaluate cosine similarity
    for c in candidates:
        if not c.embedding:
            # If embedding has not been processed yet, skip or use fallback
            c.match_score = 0.0
            continue
            
        c_vector = json.loads(c.embedding)
        sim = gemini_service.calculate_cosine_similarity(job_embedding, c_vector)
        # Scale to 0-100 percentage score
        c.match_score = round(sim * 100, 1)

    db.commit()

    # 3. Fetch and return ranked Top 10 profiles
    top_candidates = db.query(Candidate).order_by(Candidate.match_score.desc()).limit(10).all()
    return top_candidates

# ==========================================
# INTERVIEW SIMULATOR PIPELINE
# ==========================================

@app.post("/api/interviews/dispatch/{candidate_id}")
def dispatch_interview(candidate_id: int, db: Session = Depends(get_db)):
    """
    Triggers mock automated screening link dispatching to candidate.
    """
    try:
        link = simulator.dispatch_screening_link(candidate_id, db)
        return {"status": "success", "invite_link": link}
    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e))

@app.get("/api/interviews/{candidate_id}", response_model=schemas.InterviewLogResponse)
def get_interview(candidate_id: int, db: Session = Depends(get_db)):
    """
    Fetches the active interview session for the candidate.
    """
    log = db.query(InterviewLog).filter(InterviewLog.candidate_id == candidate_id).first()
    if not log:
        # Create one dynamically if candidate exists but log doesn't
        candidate = db.query(Candidate).filter(Candidate.id == candidate_id).first()
        if not candidate:
            raise HTTPException(status_code=404, detail="Candidate not found.")
        questions = simulator.generate_default_questions(candidate.skills)
        log = InterviewLog(
            candidate_id=candidate_id,
            question_1=questions[0],
            question_2=questions[1],
            question_3=questions[2],
            status="Pending"
        )
        db.add(log)
        db.commit()
        db.refresh(log)
    return log

@app.post("/api/interviews/{candidate_id}", response_model=schemas.InterviewLogResponse)
def submit_interview(candidate_id: int, submission: schemas.InterviewLogUpdate, db: Session = Depends(get_db)):
    """
    Submits candidate answers, runs automated performance scoring index,
    and updates the database record.
    """
    log = db.query(InterviewLog).filter(InterviewLog.candidate_id == candidate_id).first()
    if not log:
        raise HTTPException(status_code=404, detail="Interview session not found.")

    log.answer_1 = submission.answer_1
    log.answer_2 = submission.answer_2
    log.answer_3 = submission.answer_3
    
    # Calculate performance score
    answers = [submission.answer_1, submission.answer_2, submission.answer_3]
    performance_score = simulator.evaluate_interview_responses(answers)
    
    log.simulated_performance_index = performance_score
    log.status = "Completed"
    log.completed_at = datetime.datetime.utcnow()
    
    db.commit()
    db.refresh(log)
    return log

@app.get("/api/interviews", response_model=List[schemas.InterviewLogResponse])
def list_interviews(db: Session = Depends(get_db)):
    """
    Lists all simulated interviews logs.
    """
    return db.query(InterviewLog).order_by(InterviewLog.completed_at.desc()).all()

# ==========================================
# SYSTEM METRICS & STATS
# ==========================================

@app.get("/api/stats", response_model=schemas.DashboardStats)
def get_dashboard_stats(db: Session = Depends(get_db)):
    """
    Aggregates candidate dashboard stats cards metrics.
    """
    total_candidates = db.query(Candidate).count()
    total_anomalies = db.query(Candidate).filter(Candidate.is_anomaly == True).count()
    
    interviews = db.query(InterviewLog).filter(InterviewLog.status == "Completed").all()
    completed_interviews = len(interviews)
    
    avg_performance = 0.0
    if completed_interviews > 0:
        avg_performance = sum(i.simulated_performance_index for i in interviews) / completed_interviews
        
    return {
        "total_candidates": total_candidates,
        "total_anomalies": total_anomalies,
        "completed_interviews": completed_interviews,
        "average_performance_index": round(avg_performance, 1)
    }

# ==========================================
# RUNNER WEBHOOKS & AUTOSCALING HOOKS
# ==========================================

@app.get("/api/hooks/validate/{candidate_id}")
def validate_candidate_hook(candidate_id: int, db: Session = Depends(get_db)):
    """
    Validation validation state callback for external script runners like Hugging Face ml-intern.
    """
    candidate = db.query(Candidate).filter(Candidate.id == candidate_id).first()
    if not candidate:
        raise HTTPException(status_code=404, detail="Candidate not found.")
        
    return {
        "candidate_id": candidate.id,
        "name": candidate.name,
        "skills_density": len(candidate.skills.split(",")) / max(candidate.experience_years, 0.5),
        "z_score": candidate.z_score,
        "is_anomaly": candidate.is_anomaly,
        "persona": candidate.resume_tone,
        "validation_state": "VERIFIED_ANOMALY" if candidate.is_anomaly else "PASSED"
    }

@app.post("/api/hooks/scale")
def run_dynamic_scale_hook(db: Session = Depends(get_db)):
    """
    Evaluates dynamic resource scaling handlers based on current candidate database size.
    """
    candidates_count = db.query(Candidate).count()
    scale_metrics = hooks.scaling_handler.evaluate_scale(candidates_count)
    return {
        "status": "success",
        "hook_receiver": "ml-intern",
        "scale_assessment": scale_metrics
    }

@app.post("/api/hooks/callback")
def external_runner_callback(data: Dict):
    """
    Callback receiver for Hugging Face ml-intern evaluation loops.
    """
    logger = logging.getLogger("WebhookCallback")
    logger.info(f"Received feedback event from external runner: {data}")
    return {"status": "accepted", "message": "Callback loop evaluated successfully."}


# ==========================================
# FRONTEND MOUNT & FALLBACKS
# ==========================================

# Mount SPA static files after APIs are registered
frontend_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "frontend")
if os.path.exists(frontend_dir):
    app.mount("/", StaticFiles(directory=frontend_dir, html=True), name="frontend")
else:
    @app.get("/")
    def index_fallback():
        return {"message": "API is online, but frontend assets are missing."}
